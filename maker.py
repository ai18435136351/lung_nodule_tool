from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import time


def generate():
    document = Document()
    document.add_heading('病人病历报告', 0)

    # time tool
    time_tup = time.localtime(time.time())
    # print(time_tup)
    format_time = '%Y-%m-%d %a %H:%M:%S'
    cur_time = time.strftime(format_time, time_tup)
    # print(cur_time)

    df = pd.read_csv('information.csv')
    name = df['Name'].values[0]
    volume = df['Volume'].values[0]
    mass = df['Mass'].values[0]
    sex = df['Sex'].values[0]
    age = df['Age'].values[0]
    # print(name, volume, mass)

    gf = pd.read_csv('image_information.csv')
    image_raw = gf['image_raw'].values
    image_masked = gf['image_masked'].values
    image_value = gf['value'].values
    # print(image_raw, image_raw, image_value)

    p = document.add_paragraph(cur_time)
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    document.add_heading('病人信息', level=1)
    # document.add_paragraph('肺结节诊断', style='Intense Quote')

    document.add_paragraph('姓名:', style='List Bullet')
    document.add_paragraph('        '+name.capitalize())
    document.add_paragraph('性别:', style='List Bullet')
    document.add_paragraph('        '+sex.capitalize())
    document.add_paragraph('年龄:', style='List Bullet')
    document.add_paragraph('        '+str(age))

    # document.add_paragraph(
    #     'first item in ordered list', style='List Number'
    # )

    records = (
        (name, volume, mass)
    )
    # print(records)

    document.add_heading('肺结节情况', level=1)
    p = document.add_paragraph()
    r = p.add_run()
    for i in range(len(image_masked)):
        r.add_picture(image_raw[i], width=Inches(1.25))
        r.add_text('                                     ')
        r.add_picture(image_masked[i], width=Inches(1.25))
        r.add_text('                                     ')
        r.add_text(str(image_value[i]))

    table = document.add_table(rows=1, cols=3, style="Table Grid")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Volume'
    hdr_cells[2].text = 'Mass'
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = str(volume)
    row_cells[2].text = str(mass)

    document.add_page_break()
    document.save('病历报告.docx')
